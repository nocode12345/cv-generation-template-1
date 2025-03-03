from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
from datetime import datetime
import logging

# Configure logging for debugging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

app = Flask(__name__)

def add_line(doc, width_percent, height_pt, color, spacing_before_pt, spacing_after_pt):
    """Add a horizontal line to the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    line = OxmlElement('w:pict')
    shape = OxmlElement('v:rect')
    shape.set(qn('v:style'), 'solid')
    width_value = float(width_percent.strip('%')) / 100 * 914400  # Width in EMUs (1 cm = 914400 EMUs)
    height_value = Pt(float(height_pt.strip('pt'))).pt * 12700  # Height in EMUs (1 pt = 12700 EMUs)
    shape.set(qn('v:width'), str(int(width_value)))
    shape.set(qn('v:height'), str(int(height_value)))
    validated_color = validate_color(color)
    fill = OxmlElement('v:fill')
    fill.set(qn('color2'), validated_color[1:])
    shape.append(fill)
    line.append(shape)
    paragraph._p.append(line)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(spacing_before_pt)
    paragraph.paragraph_format.space_after = Pt(spacing_after_pt)

def validate_color(color):
    """Validate and normalize a hexadecimal color code (e.g., '#000000' or '#FFF')."""
    if not isinstance(color, str):
        logger.warning(f"Invalid color type for value: {color}, defaulting to #000000")
        return "#000000"
    
    color = color.strip().lstrip('#')
    if not color:
        logger.warning(f"Empty color value, defaulting to #000000")
        return "#000000"
    
    try:
        if len(color) not in (3, 6):
            logger.warning(f"Invalid hex color length: {color} (must be 3 or 6 chars after '#'), defaulting to #000000")
            return "#000000"
        int(color, 16)  # Test if it's a valid hex number
        if len(color) == 3:
            color = ''.join(c * 2 for c in color)
        return f"#{color}"
    except ValueError as e:
        logger.warning(f"Invalid hex color: {color} - {str(e)}, defaulting to #000000")
        return "#000000"

@app.route('/generate_cv', methods=['POST'])
def generate_cv():
    logger.debug("Received request for /generate_cv")
    try:
        # Validate and parse JSON
        if not request.is_json:
            logger.error("No JSON data received")
            return jsonify({"error": "Invalid JSON data"}), 400
        data = request.json
        logger.debug(f"JSON data received: {json.dumps(data, indent=2)}")

        # Create Word document (no required sections check)
        doc = Document()

        # Set A4 page size and margins (21cm x 29.7cm, 2.54cm margins)
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        # Header: Personal Information and Contact Details (optional)
        personal_info = data.get('personalInformation', {})
        contact_details = data.get('contactDetails', {})

        # Name (Heading 1: Arial, 20pt, bold, centered) - optional
        name = personal_info.get('name', 'N/A')
        if name != "N/A":
            p = doc.add_paragraph(name, style='Heading 1')
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(20)
            p.runs[0].bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)

        # Contact Details (Arial, 10pt, centered) - optional
        contact_text = []
        if contact_details.get('phone', 'N/A') != "N/A":
            contact_text.append(contact_details['phone'])
        if contact_details.get('email', 'N/A') != "N/A":
            contact_text.append(contact_details['email'])
        if contact_details.get('website', 'N/A') != "N/A":
            contact_text.append(contact_details['website'])
        location = contact_details.get('location', {})
        if location.get('city', 'N/A') != "N/A" and location.get('countryCode', 'N/A') != "N/A":
            contact_text.append(f"{location['city']}, {location['countryCode']}")
        if contact_text:
            contact_str = " | ".join(contact_text)
            p = doc.add_paragraph(contact_str)
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)

        # Overview: Desired Role and Tagline (Arial, 14pt, bold, centered; 10pt for tagline) - optional
        overview = data.get('overview', {})
        if overview.get('desired_role', 'N/A') != "N/A":
            p = doc.add_paragraph(overview['desired_role'])
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(14)
            p.runs[0].bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
        if overview.get('tagline', 'N/A') != "N/A":
            p = doc.add_paragraph(overview['tagline'])
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)

        # Add line after header (if header exists)
        if name != "N/A" or contact_text:
            add_line(doc, '100%', '1pt', '#000000', 12, 6)

        # Process Overview Sections (e.g., Professional Overview, Career Highlights) - optional
        for key, value in overview.items():
            if key not in ['desired_role', 'tagline'] and value and (isinstance(value, list) and value and value[0] != "N/A"):
                # Add heading (Arial, 14pt, bold)
                doc.add_heading(key, level=2).runs[0].font.name = 'Arial'
                doc.paragraphs[-1].runs[0].font.size = Pt(14)
                doc.paragraphs[-1].runs[0].bold = True
                doc.paragraphs[-1].paragraph_format.space_before = Pt(12)
                doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

                # Add bullets or paragraphs (Arial, 10pt)
                for item in value:
                    if item.startswith('•') or item.startswith(''):
                        p = doc.add_paragraph(item, style='List Bullet')
                        p.runs[0].font.name = 'Arial'
                        p.runs[0].font.size = Pt(10)
                        p.paragraph_format.left_indent = Cm(0.63)
                        p.paragraph_format.line_spacing = 1.15
                    else:
                        p = doc.add_paragraph(item)
                        p.runs[0].font.name = 'Arial'
                        p.runs[0].font.size = Pt(10)
                        p.paragraph_format.line_spacing = 1.15
                doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

        # Work Experience (Arial, 12pt, bold for job titles; 10pt for details) - optional
        work_experience = data.get('workExperience', [])
        if work_experience:
            doc.add_heading('Career Experience', level=2).runs[0].font.name = 'Arial'
            doc.paragraphs[-1].runs[0].font.size = Pt(14)
            doc.paragraphs[-1].runs[0].bold = True
            doc.paragraphs[-1].paragraph_format.space_before = Pt(12)
            doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

            for exp in work_experience:
                for pos in exp.get('position', []):
                    role_text = f"{pos.get('jobTitle', 'N/A')} - {exp.get('organisation', 'N/A')} | {exp.get('about_the_organisation', 'N/A')}, {exp.get('location', 'N/A')} ({pos.get('startDate', 'N/A')} – {pos.get('endDate', 'N/A') if pos.get('endDate', 'N/A') != 'N/A' else 'Present'})"
                    p = doc.add_paragraph(role_text)
                    p.runs[0].font.name = 'Arial'
                    p.runs[0].font.size = Pt(12)
                    p.runs[0].bold = True
                    p.paragraph_format.space_after = Pt(6)

                    # Plain text (Arial, 10pt)
                    plain_text = pos['details'].get('plainText', 'N/A') if 'details' in pos else 'N/A'
                    if plain_text != "N/A":
                        if isinstance(plain_text, list):
                            for item in plain_text:
                                p = doc.add_paragraph(item)
                                p.runs[0].font.name = 'Arial'
                                p.runs[0].font.size = Pt(10)
                                p.paragraph_format.line_spacing = 1.15
                        else:
                            p = doc.add_paragraph(plain_text)
                            p.runs[0].font.name = 'Arial'
                            p.runs[0].font.size = Pt(10)
                            p.paragraph_format.line_spacing = 1.15

                    # Key contributions or other subsections (Arial, 10pt, bullets)
                    details = pos.get('details', {})
                    for subkey, subvalue in details.items():
                        if subkey not in ['plainText'] and subvalue and isinstance(subvalue, list) and subvalue[0] != "N/A":
                            p = doc.add_paragraph(subkey)
                            p.runs[0].font.name = 'Arial'
                            p.runs[0].font.size = Pt(10)
                            p.runs[0].bold = True
                            p.paragraph_format.space_after = Pt(6)
                            for item in subvalue:
                                p = doc.add_paragraph(item, style='List Bullet')
                                p.runs[0].font.name = 'Arial'
                                p.runs[0].font.size = Pt(10)
                                p.paragraph_format.left_indent = Cm(0.63)
                                p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(6)

        # Education (Arial, 14pt, bold for heading; 10pt for details) - optional
        education = data.get('education', [])
        if education:
            doc.add_heading('Education', level=2).runs[0].font.name = 'Arial'
            doc.paragraphs[-1].runs[0].font.size = Pt(14)
            doc.paragraphs[-1].runs[0].bold = True
            doc.paragraphs[-1].paragraph_format.space_before = Pt(12)
            doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

            for edu in education:
                edu_text = f"{edu.get('studyType', 'N/A')} in {edu.get('area', 'N/A')} - {edu.get('institution', 'N/A')}, {edu.get('location', 'N/A')} ({edu.get('score', 'N/A') if edu.get('score', 'N/A') != 'N/A' else ''})"
                p = doc.add_paragraph(edu_text)
                p.runs[0].font.name = 'Arial'
                p.runs[0].font.size = Pt(10)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(6)

        # Skills (Arial, 14pt, bold for heading; 10pt for table) - optional
        skills = data.get('skills', [])
        if skills:
            doc.add_heading('Key Skills & Expertise', level=2).runs[0].font.name = 'Arial'
            doc.paragraphs[-1].runs[0].font.size = Pt(14)
            doc.paragraphs[-1].runs[0].bold = True
            doc.paragraphs[-1].paragraph_format.space_before = Pt(12)
            doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

            # Create a three-column table for skills
            columns = [
                {"text": "\n".join([s.get('name', 'N/A') for s in skills[:5]]), "font": {'name': 'Arial', 'size': '10pt', 'color': '#000000'}},
                {"text": "\n".join([s.get('name', 'N/A') for s in skills[5:10]]), "font": {'name': 'Arial', 'size': '10pt', 'color': '#000000'}},
                {"text": "\n".join([s.get('name', 'N/A') for s in skills[10:]]), "font": {'name': 'Arial', 'size': '10pt', 'color': '#000000'}}
            ]
            table = doc.add_table(rows=1, cols=3)
            table.autofit = False
            for col in table.columns:
                col.width = Cm(6.5)  # Roughly equal width for A4 (21cm / 3 - margins)
            for i, col_data in enumerate(columns):
                cell = table.cell(0, i)
                for text in col_data['text'].split('\n'):
                    if text.strip() and text != "N/A":
                        p = cell.add_paragraph(text)
                        p.runs[0].font.name = 'Arial'
                        p.runs[0].font.size = Pt(10)
                        p.paragraph_format.line_spacing = 1.15
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT
            table.paragraphs[0].paragraph_format.space_after = Pt(6)

        # Other sections (associations, publications, projects, volunteer, interests, patents, awards, certificates, languages, references) - optional
        for section in ['associations', 'publications', 'projects', 'volunteer', 'interests', 'patents', 'awards', 'certificates', 'languages', 'references']:
            section_data = data.get(section)
            if section_data and (isinstance(section_data, list) and section_data and section_data[0] != "N/A" or isinstance(section_data, str) and section_data != "N/A"):
                doc.add_heading(section.capitalize(), level=2).runs[0].font.name = 'Arial'
                doc.paragraphs[-1].runs[0].font.size = Pt(14)
                doc.paragraphs[-1].runs[0].bold = True
                doc.paragraphs[-1].paragraph_format.space_before = Pt(12)
                doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

                if isinstance(section_data, list):
                    for item in section_data:
                        if isinstance(item, dict):
                            text = f"{item.get('name', 'N/A')} - {item.get('summary', '') if 'summary' in item else ''}"
                            if text.strip() != "N/A - ":
                                p = doc.add_paragraph(text)
                                p.runs[0].font.name = 'Arial'
                                p.runs[0].font.size = Pt(10)
                                p.paragraph_format.line_spacing = 1.15
                                p.paragraph_format.space_after = Pt(6)
                        else:
                            p = doc.add_paragraph(item, style='List Bullet')
                            p.runs[0].font.name = 'Arial'
                            p.runs[0].font.size = Pt(10)
                            p.paragraph_format.left_indent = Cm(0.63)
                            p.paragraph_format.line_spacing = 1.15
                else:
                    p = doc.add_paragraph(section_data)
                    p.runs[0].font.name = 'Arial'
                    p.runs[0].font.size = Pt(10)
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(6)

        # Add line after each major section (except the last one)
        for i in range(len(doc.paragraphs) - 1):
            if doc.paragraphs[i].style.name == 'Heading 2':
                add_line(doc, '100%', '1pt', '#000000', 12, 6)

        # Save document to temporary directory (Render-safe)
        output_filename = f"CV_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join('/tmp', output_filename)
        os.makedirs('/tmp', exist_ok=True)
        doc.save(output_path)
        logger.debug(f"Word document generated and saved at {output_path}")
        return send_file(output_path, as_attachment=True, download_name=output_filename, 
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except KeyError as e:
        logger.error(f"KeyError in JSON processing: {str(e)}")
        return jsonify({"error": f"Missing key in JSON: {str(e)}"}), 400
    except ValueError as e:
        logger.error(f"ValueError in JSON processing: {str(e)}")
        return jsonify({"error": f"Invalid value in JSON: {str(e)}"}), 400
    except Exception as e:
        logger.error(f"Error in /generate_cv: {str(e)}")
        return jsonify({"error": f"Internal Server Error: {str(e)}"}), 500

if __name__ == '__main__':
    # Use Render's PORT (no fallback) for local testing
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
